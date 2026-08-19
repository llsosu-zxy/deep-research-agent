from __future__ import annotations

import argparse
import re
from pathlib import Path

import docx

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = Path(r"C:\Users\29716\Documents\实习\新加坡AI实习机会清单.docx")
OUTPUT_DIR = ROOT / "data" / "corpus" / "imported"

HEADER_MAP = {
    "公司 / 项目": "company",
    "岗位方向": "roles",
    "薪资待遇": "compensation",
    "能否转正": "return_offer",
    "笔试 / 面试": "hiring_process",
    "官方申请链接": "apply_url",
    "备注 / 推荐理由": "notes",
}


def _slugify(text: str, max_len: int = 60) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
    return cleaned[:max_len]


def _front_matter(company: str, roles: str, url: str, tags: list[str]) -> str:
    tag_text = ", ".join(tags[:8])
    return (
        "---\n"
        f"title: {company} - {roles[:80]}\n"
        f"source_url: {url}\n"
        "source_type: job_page\n"
        "collected_at: 2026-08-14\n"
        f"tags: [{tag_text}]\n"
        "---\n"
    )


def _render_doc(record: dict) -> str:
    lines = [
        f"# {record['company']} AI / Data Internship",
        "",
        "## Roles",
        record["roles"],
        "",
        "## Compensation",
        record["compensation"],
        "",
        "## Return Offer",
        record["return_offer"],
        "",
        "## Hiring Process",
        record["hiring_process"],
        "",
        "## Notes",
        record["notes"],
    ]
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Import the Singapore AI internship docx into the corpus.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT_DIR)
    args = parser.parse_args()

    document = docx.Document(str(args.source))
    if not document.tables:
        raise RuntimeError("Source docx contains no table")
    table = document.tables[0]
    rows = list(table.rows)
    header = [cell.text.strip() for cell in rows[0].cells]
    records = []
    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        record = {}
        for header_name, cell in zip(header, cells):
            key = HEADER_MAP.get(header_name, header_name)
            record[key] = cell
        company = record.get("company", "").strip()
        roles = record.get("roles", "").strip()
        if not company or not roles:
            continue
        records.append(record)

    args.output.mkdir(parents=True, exist_ok=True)
    master_lines = [
        "# Singapore AI / Data Internship Master List",
        "",
        "Imported from the internship opportunity sheet on 2026-08-14.",
        "",
        "| Company | Roles | Compensation | Return | Process | Apply |",
        "|---|---|---|---|---|---|",
    ]
    for index, record in enumerate(records, start=1):
        company = record["company"]
        roles = record["roles"]
        url = record.get("apply_url", "").strip()
        tags = [part.strip() for part in company.replace("/", " ").split() if part.strip()] or ["singapore", "internship"]
        filename = f"{index:02d}_{_slugify(company)}.md"
        front_matter = _front_matter(company, roles, url, tags)
        (args.output / filename).write_text(front_matter + _render_doc(record).strip() + "\n", encoding="utf-8")
        master_lines.append(
            f"| {company} | {roles[:80]} | {record.get('compensation', '')[:40]} | "
            f"{record.get('return_offer', '')[:40]} | {record.get('hiring_process', '')[:60]} | {url} |"
        )
    (args.output / "00_master-list.md").write_text("\n".join(master_lines) + "\n", encoding="utf-8")
    print(f"Imported {len(records)} records into {args.output}")


if __name__ == "__main__":
    main()
